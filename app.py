import gradio as gr
import os
import json
import docx
from docx.shared import RGBColor
from dotenv import load_dotenv

# LangChain Imports
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain.chains.retrieval_qa import RetrievalQA

# --- 1. CONFIGURATION & INITIALIZATION ---
load_dotenv()

# Check for API Key
if not os.getenv("GOOGLE_API_KEY"):
    raise ValueError("GOOGLE_API_KEY not found in .env file. Please set it.")

# Check for FAISS index
if not os.path.exists("faiss_index"):
    raise FileNotFoundError("FAISS index not found. Please run 'create_vector_store.py' first.")

# Initialize LLM and Embeddings
llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro-latest", temperature=0.2, convert_system_message_to_human=True)
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

# Load the FAISS vector store
db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)

# --- 2. DOCUMENT CHECKLIST & PROCESS IDENTIFICATION ---
DOCUMENT_CHECKLISTS = {
    "Company Incorporation": {
        "required_count": 5,
        "documents": {
            "articles of association": "Articles of Association",
            "memorandum of association": "Memorandum of Association (MoA/MoU)",
            "incorporation application form": "Incorporation Application Form",
            "ubo declaration form": "UBO Declaration Form",
            "register of members and directors": "Register of Members and Directors"
        }
    }
}

def identify_process_and_missing_docs(file_paths):
    uploaded_docs = [os.path.basename(f.name).lower() for f in file_paths]
    process = "Company Incorporation"
    checklist = DOCUMENT_CHECKLISTS[process]
    
    found_docs = {official_name for doc_name in uploaded_docs for keyword, official_name in checklist["documents"].items() if keyword in doc_name}
    
    missing_docs = [doc for doc in checklist["documents"].values() if doc not in found_docs]
    
    return {
        "process": process,
        "documents_uploaded_count": len(found_docs),
        "required_documents_count": checklist["required_count"],
        "missing_documents": missing_docs
    }

# --- 3. DOCUMENT HIGHLIGHTING & COMMENTING ---
def highlight_and_comment(doc, issue):
    """Adds highlights and comments to the docx object for a single issue."""
    text_to_find = issue.get("offending_text", "")
    if not text_to_find:
        return # Cannot highlight if we don't know what text to look for

    comment_text = f"Issue: {issue.get('issue', 'N/A')}\nSuggestion: {issue.get('suggestion', 'N/A')}"
    
    # Iterate through paragraphs to find and highlight text
    for para in doc.paragraphs:
        if text_to_find.lower() in para.text.lower():
            # Highlighting is done by splitting a paragraph into runs
            # This is a simplified implementation. Complex formatting might be lost.
            inline = para.runs
            # Replace the paragraph's text with a new set of runs
            para.clear()
            # Split text based on the found issue text (case-insensitive)
            parts = para.text.split(text_to_find)
            for i, part in enumerate(parts):
                para.add_run(part)
                if i < len(parts) - 1:
                    # Add the highlighted run
                    highlighted_run = para.add_run(text_to_find)
                    highlighted_run.font.color.rgb = RGBColor(255, 0, 0) # Red color
                    highlighted_run.bold = True
            
            # Add the comment to the paragraph
            para.add_comment(comment_text, author="Corporate Agent")
            return # Stop after finding the first occurrence in a paragraph

def create_reviewed_docx(original_file, issues_found):
    """Creates a new .docx file with highlights and comments."""
    if not issues_found or "Critical" in {issue.get("severity") for issue in issues_found}:
        return None

    doc = docx.Document(original_file.name)
    
    for issue in issues_found:
        highlight_and_comment(doc, issue)

    base, ext = os.path.splitext(os.path.basename(original_file.name))
    reviewed_filename = f"Reviewed_{base}.docx"
    temp_dir = "temp_reviewed_docs"
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)
    reviewed_filepath = os.path.join(temp_dir, reviewed_filename)
    
    doc.save(reviewed_filepath)
    return reviewed_filepath

# --- 4. LANGCHAIN RAG ANALYSIS ---
def analyze_document_with_langchain(doc_text, doc_name):
    """Analyzes a document using a LangChain RetrievalQA chain."""
    
    # Define the prompt template for the final analysis
    prompt_template = """
    You are an AI legal assistant specializing in Abu Dhabi Global Market (ADGM) regulations.
    Your task is to review the legal document named '{doc_name}' based on the provided context.

    **Instructions:**
    1. Use the **Provided ADGM Legal Context** below to ensure your analysis is accurate.
    2. Identify legal red flags, missing clauses, and non-compliance issues.
    3. For each issue, you MUST identify the specific text from the document that is problematic.
    4. Return your findings ONLY in a structured JSON format. The JSON should be a list of objects. Do not add any text or explanation outside the JSON structure.

    **JSON Output Format for each issue:**
    {{
      "section": "The clause or section number, e.g., 'Clause 3.1'",
      "offending_text": "The exact text from the document that contains the issue.",
      "issue": "A clear, one-sentence description of the issue found.",
      "severity": "A severity rating: 'High', 'Medium', or 'Low'.",
      "suggestion": "A compliant suggestion or an action to be taken.",
      "citation": "The specific ADGM rule that applies, based on the context."
    }}
    
    ---
    **Provided ADGM Legal Context:**
    {context}
    ---
    **Document Text for Analysis:**
    {question}
    ---

    Now, provide your analysis as a single JSON list.
    """
    
    PROMPT = PromptTemplate(
        template=prompt_template, input_variables=["context", "question", "doc_name"]
    )
    
    # Create the RetrievalQA chain
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=db.as_retriever(search_kwargs={"k": 4}), # Retrieve top 4 relevant chunks
        return_source_documents=True,
        chain_type_kwargs={"prompt": PROMPT.partial(doc_name=doc_name)}
    )

    # Run the chain
    try:
        result = qa_chain.invoke({"query": doc_text})
        # The result from the LLM is in the 'result' key
        # We need to find the JSON part of the string
        json_string = result['result'].strip()
        json_start = json_string.find('[')
        json_end = json_string.rfind(']') + 1
        if json_start != -1 and json_end != -1:
            clean_json = json_string[json_start:json_end]
            return json.loads(clean_json)
        else:
            return [{"issue": "Could not parse LLM response.", "severity": "Critical"}]

    except Exception as e:
        print(f"Error during LangChain analysis: {e}")
        return [{"issue": f"Failed to analyze document due to an internal error: {e}", "severity": "Critical"}]


# --- 5. MAIN GRADIO PROCESSING FUNCTION ---
def process_documents(files):
    if files is None:
        return "Please upload documents to begin.", None, None
        
    checklist_result = identify_process_and_missing_docs(files)
    
    if checklist_result['missing_documents']:
        missing_docs_str = ", ".join(checklist_result['missing_documents'])
        checklist_notification = (
            f"It appears that you're trying to {checklist_result['process']}. "
            f"Based on our reference list, you have uploaded {checklist_result['documents_uploaded_count']} out of "
            f"{checklist_result['required_documents_count']} required documents. "
            f"The missing document(s) appear to be: '{missing_docs_str}'."
        )
    else:
        checklist_notification = f"Document checklist passed for {checklist_result['process']}. All required documents seem to be present."

    all_issues = []
    reviewed_file_paths = []
    
    for file_obj in files:
        doc_name = os.path.basename(file_obj.name)
        try:
            doc = docx.Document(file_obj.name)
            full_text = "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            all_issues.append({"document": doc_name, "issue": f"Could not read .docx file: {e}", "severity": "Critical"})
            continue
            
        issues = analyze_document_with_langchain(full_text, doc_name)
        
        if issues:
            for issue in issues:
                issue['document'] = doc_name
            all_issues.extend(issues)
            
            reviewed_path = create_reviewed_docx(file_obj, issues)
            if reviewed_path:
                reviewed_file_paths.append(reviewed_path)

    final_report = {
        "process": checklist_result["process"],
        "documents_uploaded": checklist_result["documents_uploaded_count"],
        "required_documents": checklist_result["required_documents_count"],
        "missing_document(s)": checklist_result["missing_documents"],
        "issues_found": all_issues
    }
    
    return checklist_notification, final_report, reviewed_file_paths

# --- 6. GRADIO UI ---
with gr.Blocks(theme=gr.themes.Soft()) as demo:
    gr.Markdown("# 🤖 ADGM-Compliant Corporate Agent (LangChain & FAISS)")
    gr.Markdown("This AI assistant uses a RAG architecture to analyze legal documents against ADGM regulations, providing comments and highlights directly in the file.")
    gr.Markdown("---")
    
    with gr.Row():
        with gr.Column(scale=1):
            gr.Markdown("### 1. Upload Documents for Analysis")
            file_uploads = gr.File(
                file_count="multiple",
                file_types=[".docx"],
                label="Upload company formation documents (.docx)"
            )
            submit_btn = gr.Button("Analyze Documents", variant="primary")

        with gr.Column(scale=2):
            gr.Markdown("### 2. Analysis & Results")
            checklist_output = gr.Textbox(label="Checklist Verification Status", lines=3, interactive=False)
            json_output = gr.JSON(label="Structured Analysis Report")
            reviewed_files_output = gr.File(label="Download Reviewed Documents (with Highlights & Comments)", file_count="multiple", interactive=False)

    submit_btn.click(
        fn=process_documents,
        inputs=file_uploads,
        outputs=[checklist_output, json_output, reviewed_files_output]
    )
    
    gr.Markdown("---")
    gr.Markdown("⚠️ **Important:** Before starting, ensure you have run `python create_vector_store.py` to build the knowledge base.")

if __name__ == "__main__":
    demo.launch()
